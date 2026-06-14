import traceback
try:
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    doc.add_heading('Test', 1)
    doc.add_paragraph('Hello World')
    output = r'c:\Users\ADVAN\source\repos\cafe_yo\cafe_yo\docs\test_output.docx'
    doc.save(output)
    import os
    print(f'EXISTS: {os.path.exists(output)}, SIZE: {os.path.getsize(output)}')
except Exception as e:
    traceback.print_exc()
    print(f'ERROR: {e}')
